import io
import re
import logging
import asyncio
from datetime import datetime
from typing import List, Dict, TypedDict, Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# File parsing libraries
import PyPDF2
import docx

# LangChain & AI Components
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangGraph for workflow orchestration
from langgraph.graph import StateGraph, END

# --- Configuration & Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Intelligent Resume Ranker API V3 (Gemini Edition)",
    description="A LangGraph-powered, two-layer resume screening system with Google Gemini embeddings and Groq generative AI feedback.",
    version="3.2.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)

# --- Pydantic Models for API ---
class RankedResume(BaseModel):
    rank: int
    filename: str
    semantic_score: float
    skill_match_count: int
    total_skills_required: int
    matched_skills: List[str]
    missing_skills: List[str]
    content_preview: str
    key_matching_chunk: str
    ai_verdict: Optional[str] = "N/A"
    ai_feedback: Optional[str] = "N/A"

class RankingResult(BaseModel):
    job_title: str
    total_resumes_processed: int
    top_candidates: List[RankedResume]
    processing_time: float

# --- Core Text Processing ---
class TextProcessor:
    def clean_text(self, text: str) -> str:
        text = re.sub(r'[^\w\s\.,\-\+#@]', ' ', text)
        return re.sub(r'\s+', ' ', text).strip()

    def extract_text_from_file(self, content: bytes, filename: str) -> str:
        try:
            ext = filename.split('.')[-1].lower()
            if ext == 'pdf':
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == 'docx':
                doc = docx.Document(io.BytesIO(content))
                return "\n".join(para.text for para in doc.paragraphs)
            elif ext == 'txt':
                return content.decode('utf-8', errors='ignore')
            return ""
        except Exception as e:
            logger.error(f"Failed to process {filename}: {e}")
            return ""

    def extract_sections(self, text: str) -> Dict[str, str]:
        sections, patterns = {}, {
            'experience': r'(experience|employment history|work history|professional experience)\s*:?',
            'skills': r'(skills|technical skills|core competencies|technologies)\s*:?'
        }
        lower_text = text.lower()
        starts = {sec: m.start() for sec, pat in patterns.items() if (m := re.search(pat, lower_text, re.IGNORECASE))}
        sorted_secs = sorted(starts.items(), key=lambda i: i[1])
        for i, (sec, start) in enumerate(sorted_secs):
            end = sorted_secs[i+1][1] if i + 1 < len(sorted_secs) else len(text)
            content = text[start:end].strip()
            header = re.match(patterns[sec], content, re.IGNORECASE | re.DOTALL)
            if header: content = content[header.end():].strip()
            if content: sections[sec] = content
        return sections

# --- LangGraph State Definition ---
class AgentState(TypedDict):
    job_description: str
    required_skills: List[str]
    file_data_list: List[tuple]
    top_n: int
    groq_api_key: str
    gemini_api_key: str
    section_documents: List[Document]
    processed_resumes: Dict[str, Dict]
    semantic_candidates: List[Dict]
    ranked_candidates: List[Dict]

# --- LangGraph Nodes ---
class ResumeWorkflow:
    def __init__(self):
        self.text_processor = TextProcessor()

    def preprocess_resumes(self, state: AgentState) -> AgentState:
        processed, raw_docs = {}, []
        for fname, content in state["file_data_list"]:
            full_text = self.text_processor.extract_text_from_file(content, fname)
            if not full_text: continue
            sections = self.text_processor.extract_sections(full_text)
            processed[fname] = {'full_text': full_text, 'sections': sections}
            for sec, text in sections.items():
                raw_docs.append(Document(page_content=text, metadata={'filename': fname, 'section': sec}))
            if not sections:
                raw_docs.append(Document(page_content=full_text, metadata={'filename': fname, 'section': 'general'}))
        
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(raw_docs)
        logger.info(f"Split {len(raw_docs)} sections into {len(chunks)} chunks.")

        state["processed_resumes"] = processed
        state["section_documents"] = chunks
        return state

    def semantic_search(self, state: AgentState) -> AgentState:
        if not state["section_documents"]:
            state["semantic_candidates"] = []
            return state
        try:
            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=state["gemini_api_key"])
        except Exception as e:
            raise RuntimeError(f"Gemini Embeddings initialization failed: {e}")

        vector_store = Chroma.from_documents(state["section_documents"], embeddings)
        similar = vector_store.similarity_search_with_relevance_scores(state["job_description"], k=100)
        
        candidates = {}
        for doc, score in similar:
            fname = doc.metadata['filename']
            if fname not in candidates or score > candidates[fname]['score']:
                candidates[fname] = {
                    'filename': fname, 'score': score,
                    'key_matching_chunk': f"[{doc.metadata.get('section', 'general').upper()}] {doc.page_content[:500]}..."
                }
        vector_store.delete_collection()
        state["semantic_candidates"] = list(candidates.values())
        return state

    def skill_rerank(self, state: AgentState) -> AgentState:
        reranked = []
        for cand in state["semantic_candidates"]:
            fname = cand['filename']
            full_text = state["processed_resumes"][fname]['full_text'].lower()
            matched = [s for s in state["required_skills"] if re.search(r'(?<!\w)' + re.escape(s.strip().lower()) + r'(?!\w)', full_text)]
            reranked.append({**cand, 'skill_match_count': len(matched), 'matched_skills': matched, 'missing_skills': [s for s in state["required_skills"] if s not in matched]})
        reranked.sort(key=lambda x: (-x['skill_match_count'], -x['score']))
        state["ranked_candidates"] = reranked
        return state

    async def generate_feedback(self, state: AgentState) -> AgentState:
        if not state["groq_api_key"]:
            return state
        try:
            llm = ChatGroq(temperature=0, model_name="gemma2-9b-it", groq_api_key=state["groq_api_key"])
            parser = JsonOutputParser()
            prompt = ChatPromptTemplate.from_messages([
                ("system", 'You are an HR analyst. Provide a concise verdict and feedback on a candidate\'s fit based on their resume. Respond ONLY with a JSON object: {"ai_verdict": "Verdict", "ai_feedback": "Feedback"}.'),
                ("human", "Analyze the candidate for a role requiring: {required_skills}.\nExperience: {experience_section}\nSkills: {skills_section}\nProvide your analysis in the specified JSON format."),
            ])
            chain = prompt | llm | parser
        except Exception as e:
            logger.error(f"Groq model initialization failed: {e}. Skipping feedback.")
            return state

        async def get_feedback(candidate):
            fname = candidate['filename']
            sections = state["processed_resumes"][fname]['sections']
            try:
                response = await chain.ainvoke({
                    "required_skills": ", ".join(state["required_skills"]),
                    "experience_section": sections.get('experience', 'N/A'),
                    "skills_section": sections.get('skills', 'N/A')
                })
                candidate.update(response)
            except Exception as e:
                logger.error(f"Feedback generation error for {fname}: {e}")
                candidate.update({"ai_verdict": "Error", "ai_feedback": "Could not generate feedback."})
            return candidate

        tasks = [get_feedback(c) for c in state["ranked_candidates"][:state["top_n"]]]
        updated = await asyncio.gather(*tasks)
        for i, u in enumerate(updated): state["ranked_candidates"][i] = u
        return state

# --- LangGraph Setup ---
workflow = ResumeWorkflow()
graph = StateGraph(AgentState)
graph.add_node("preprocess", workflow.preprocess_resumes)
graph.add_node("semantic_search", workflow.semantic_search)
graph.add_node("skill_rerank", workflow.skill_rerank)
graph.add_node("feedback", workflow.generate_feedback)
graph.set_entry_point("preprocess")
graph.add_edge("preprocess", "semantic_search")
graph.add_edge("semantic_search", "skill_rerank")
graph.add_edge("skill_rerank", "feedback")
graph.add_edge("feedback", END)
app_graph = graph.compile()

# --- API Endpoint ---
@app.post("/rank-resumes", response_model=RankingResult)
async def rank_resumes_endpoint(
    job_title: str = Form(...), job_description: str = Form(...),
    required_skills: str = Form(...), groq_api_key: str = Form(""),
    gemini_api_key: str = Form(...), top_n: int = Form(10),
    files: List[UploadFile] = File(...)
):
    start = datetime.now()
    skills = [s.strip() for s in required_skills.split(',') if s.strip()]
    if not all([job_description, skills, gemini_api_key]):
        raise HTTPException(status_code=422, detail="Job Description, Skills, and Gemini API Key are required.")
    
    file_data = [(f.filename, await f.read()) for f in files if f.filename]
    if not file_data: raise HTTPException(status_code=400, detail="No valid resumes were provided.")
    
    state = AgentState(
        job_title=job_title, job_description=job_description, required_skills=skills,
        file_data_list=file_data, top_n=top_n, groq_api_key=groq_api_key,
        gemini_api_key=gemini_api_key
    )
    
    try:
        final_state = await app_graph.ainvoke(state)
        candidates = [RankedResume(
            rank=i + 1, filename=c['filename'], semantic_score=c['score'],
            skill_match_count=c['skill_match_count'], total_skills_required=len(skills),
            matched_skills=c['matched_skills'], missing_skills=c['missing_skills'],
            content_preview=final_state["processed_resumes"][c['filename']]['full_text'][:400] + "...",
            key_matching_chunk=c.get('key_matching_chunk', "N/A"),
            ai_verdict=c.get('ai_verdict', "N/A"), ai_feedback=c.get('ai_feedback', "N/A")
        ) for i, c in enumerate(final_state["ranked_candidates"][:top_n])]

        return RankingResult(
            job_title=job_title, total_resumes_processed=len(final_state["processed_resumes"]),
            top_candidates=candidates, processing_time=(datetime.now() - start).total_seconds()
        )
    except Exception as e:
        logger.error(f"Workflow execution failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

