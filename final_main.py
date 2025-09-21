import io
import re
import uuid
import logging
import asyncio
import os
import tempfile
from datetime import datetime
from typing import List, Dict, TypedDict, Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# File parsing libraries
import PyPDF2
import docx
from langchain_community.document_loaders import PDFMinerLoader 

# LangChain & AI Components
from langchain_core.documents import Document
# --- MODIFIED: Swapped Ollama for SentenceTransformerEmbeddings ---
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_chroma import Chroma
import chromadb
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser

# LangGraph for workflow orchestration
from langgraph.graph import StateGraph, END

# --- Configuration & Setup ---
# Note: To run this, you will need to install the libraries from requirements.txt
# pip install -r requirements.txt
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Intelligent Resume Ranker API V4 (Cloud Ready)",
    description="A LangGraph-powered resume screener using Sentence-Transformers for cloud deployment.",
    version="4.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
)

# --- Pydantic Models for API (Unchanged) ---
class RankedResume(BaseModel):
    rank: int
    filename: str
    semantic_score: float
    skill_match_count: int
    total_skills_required: int
    matched_skills: List[str]
    missing_skills: List[str]
    content_preview: str
    key_matching_chunk: str
    ai_verdict: Optional[str] = "N/A"
    ai_feedback: Optional[str] = "N/A"

class RankingResult(BaseModel):
    job_title: str
    total_resumes_processed: int
    top_candidates: List[RankedResume]
    processing_time: float

# --- Core Text Processing (Unchanged) ---
class TextProcessor:
    def clean_text(self, text: str) -> str:
        text = re.sub(r'[^\w\s\.,\-\+#@]', ' ', text)
        return re.sub(r'\s+', ' ', text).strip()

    def extract_text_from_file(self, content: bytes, filename: str) -> str:
        try:
            extension = filename.split('.')[-1].lower()
            if extension == 'pdf':
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            elif extension == 'docx':
                doc = docx.Document(io.BytesIO(content))
                return "\n".join(para.text for para in doc.paragraphs)
            elif extension == 'txt':
                return content.decode('utf-8', errors='ignore')
            return ""
        except Exception as e:
            logger.error(f"Failed to process {filename}: {e}")
            return ""

    def extract_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        patterns = {
            'experience': r'(experience|employment history|work history|professional experience)\s*:?',
            'skills': r'(skills|technical skills|core competencies|technologies)\s*:?'
        }
        text_lower = text.lower()
        section_starts = {sec: m.start() for sec, pat in patterns.items() if (m := re.search(pat, text_lower, re.IGNORECASE))}
        sorted_sections = sorted(section_starts.items(), key=lambda item: item[1])

        for i, (sec, start_index) in enumerate(sorted_sections):
            end_index = sorted_sections[i+1][1] if i + 1 < len(sorted_sections) else len(text)
            content = text[start_index:end_index].strip()
            header_match = re.match(patterns[sec], content, re.IGNORECASE | re.DOTALL)
            if header_match:
                content = content[header_match.end():].strip()
            if content:
                sections[sec] = content
        return sections

# --- LangGraph State Definition (Unchanged) ---
class AgentState(TypedDict):
    job_description: str
    required_skills: List[str]
    file_data_list: List[tuple]
    top_n: int
    groq_api_key: str
    
    section_documents: List[Document]
    processed_resumes: Dict[str, Dict]
    semantic_candidates: List[Dict]
    ranked_candidates: List[Dict]

# --- LangGraph Nodes ---
class ResumeWorkflow:
    def __init__(self):
        self.text_processor = TextProcessor()
        # --- MODIFIED: Initialize SentenceTransformerEmbeddings ---
        # This will download the model from Hugging Face hub automatically on first run.
        # 'all-mpnet-base-v2' is a robust model, great for this use case.
        logger.info("Initializing SentenceTransformerEmbeddings with 'all-mpnet-base-v2'...")
        self.embedding_model = SentenceTransformerEmbeddings(model_name="all-mpnet-base-v2")
        logger.info("✅ SentenceTransformerEmbeddings model initialized.")

    # --- All other nodes (preprocess_resumes, semantic_search, etc.) remain unchanged ---
    # They will now use the new self.embedding_model seamlessly.

    def preprocess_resumes(self, state: AgentState) -> AgentState:
        logger.info("--- Starting Node: Preprocess Resumes ---")
        processed_resumes = {}
        section_documents = []

        for filename, content in state["file_data_list"]:
            full_text = self.text_processor.extract_text_from_file(content, filename)
            if not full_text:
                logger.warning(f"Could not extract text from {filename}, skipping.")
                continue
            
            sections = self.text_processor.extract_sections(full_text)
            processed_resumes[filename] = {'full_text': full_text, 'sections': sections}

            for sec_name, sec_content in sections.items():
                section_documents.append(Document(
                    page_content=sec_content,
                    metadata={'filename': filename, 'section': sec_name}
                ))
            if not sections:
                 section_documents.append(Document(
                    page_content=full_text[:2000],
                    metadata={'filename': filename, 'section': 'general'}
                ))

        state["processed_resumes"] = processed_resumes
        state["section_documents"] = section_documents
        return state

    def semantic_search(self, state: AgentState) -> AgentState:
        logger.info("--- Starting Node: Semantic Search (Layer 1) ---")
        section_documents = state.get("section_documents", [])
        if not section_documents:
            state["semantic_candidates"] = []
            return state

        vector_store = Chroma.from_documents(section_documents, self.embedding_model)
        
        similar_docs_with_scores = vector_store.similarity_search_with_relevance_scores(
            state["job_description"], k=100
        )
        
        semantic_candidates_map = {}
        for doc, score in similar_docs_with_scores:
            filename = doc.metadata['filename']
            if filename not in semantic_candidates_map or score > semantic_candidates_map[filename]['score']:
                 semantic_candidates_map[filename] = {
                    'filename': filename,
                    'score': score,
                    'key_matching_chunk': f"[{doc.metadata.get('section', 'general').upper()}] {doc.page_content[:500]}..."
                }
        
        vector_store.delete_collection()
        logger.info(f"Found {len(semantic_candidates_map)} unique semantic candidates.")
        state["semantic_candidates"] = list(semantic_candidates_map.values())
        return state

    def skill_rerank(self, state: AgentState) -> AgentState:
        logger.info("--- Starting Node: Skill Re-ranking (Layer 2) ---")
        re_ranking_list = []
        required_skills = state["required_skills"]

        for candidate in state["semantic_candidates"]:
            filename = candidate['filename']
            full_text = state["processed_resumes"][filename]['full_text']
            
            resume_text_lower = full_text.lower()
            matched_skills = [
                skill for skill in required_skills 
                if re.search(r'(?<!\w)' + re.escape(skill.strip().lower()) + r'(?!\w)', resume_text_lower)
            ]
            
            re_ranking_list.append({
                **candidate,
                'skill_match_count': len(matched_skills),
                'matched_skills': matched_skills,
                'missing_skills': [s for s in required_skills if s not in matched_skills]
            })

        re_ranking_list.sort(key=lambda x: (-x['skill_match_count'], -x['score']))
        state["ranked_candidates"] = re_ranking_list
        return state

    async def generate_feedback(self, state: AgentState) -> AgentState:
        logger.info("--- Starting Node: Generate AI Feedback ---")
        if not state["groq_api_key"]:
            logger.warning("No Groq API key provided. Skipping AI feedback generation.")
            return state

        try:
            llm = ChatGroq(temperature=0, model_name="gemma2-9b-it", groq_api_key=state["groq_api_key"])
            parser = JsonOutputParser()
            
            prompt = ChatPromptTemplate.from_messages([
                ("system", """You are a formal HR analyst. Your task is to provide a direct, concise verdict and feedback on a candidate's fit for a role based on their resume sections and the job's required skills. Respond ONLY with a JSON object containing two keys: "ai_verdict" and "ai_feedback".
                - "ai_verdict": A short, direct verdict like "Strong Fit", "Potential Fit", or "Lacks Key Skills".
                - "ai_feedback": A 2-3 sentence formal analysis explaining the verdict based on the evidence."""),
                ("human", """Analyze the candidate for a role requiring: {required_skills}.
                **Candidate's Experience:** {experience_section}
                **Candidate's Skills:** {skills_section}
                Provide your analysis in the specified JSON format."""),
            ])
            chain = prompt | llm | parser
        except Exception as e:
            logger.error(f"Failed to initialize Groq model: {e}. Skipping feedback.")
            return state

        async def get_feedback_for_candidate(candidate):
            filename = candidate['filename']
            sections = state["processed_resumes"][filename]['sections']
            
            try:
                response = await chain.ainvoke({
                    "required_skills": ", ".join(state["required_skills"]),
                    "experience_section": sections.get('experience', 'Not specified.'),
                    "skills_section": sections.get('skills', 'Not specified.')
                })
                candidate.update(response)
            except Exception as e:
                logger.error(f"Error generating feedback for {filename}: {e}")
                candidate["ai_verdict"] = "Error"
                candidate["ai_feedback"] = "Could not generate AI feedback."
            return candidate

        top_candidates = state["ranked_candidates"][:state["top_n"]]
        tasks = [get_feedback_for_candidate(c) for c in top_candidates]
        updated_candidates = await asyncio.gather(*tasks)
        
        for i, updated_candidate in enumerate(updated_candidates):
            state["ranked_candidates"][i] = updated_candidate

        return state

# --- LangGraph Setup (Unchanged) ---
workflow_manager = ResumeWorkflow()
graph = StateGraph(AgentState)

graph.add_node("preprocess_resumes", workflow_manager.preprocess_resumes)
graph.add_node("semantic_search", workflow_manager.semantic_search)
graph.add_node("skill_rerank", workflow_manager.skill_rerank)
graph.add_node("generate_feedback", workflow_manager.generate_feedback)

graph.set_entry_point("preprocess_resumes")
graph.add_edge("preprocess_resumes", "semantic_search")
graph.add_edge("semantic_search", "skill_rerank")
graph.add_edge("skill_rerank", "generate_feedback")
graph.add_edge("generate_feedback", END)

app_graph = graph.compile()

# --- Main API Endpoint (Unchanged) ---
@app.post("/rank-resumes", response_model=RankingResult)
async def rank_resumes_endpoint(
    job_title: str = Form(...),
    job_description: str = Form(""),
    required_skills: str = Form(...),
    groq_api_key: str = Form(""),
    top_n: int = Form(10),
    files: List[UploadFile] = File(...),
    jd_file: Optional[UploadFile] = File(None)
):
    start_time = datetime.now()
    
    skill_list = [skill.strip() for skill in required_skills.split(',') if skill.strip()]

    effective_jd = job_description
    if jd_file:
        logger.info(f"Processing JD from uploaded file: {jd_file.filename}")
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(await jd_file.read())
                tmp_path = tmp.name

            loader = PDFMinerLoader(tmp_path)
            documents = loader.load()
            effective_jd = " ".join([doc.page_content for doc in documents])
            logger.info("Successfully extracted text from JD PDF.")

        except Exception as e:
            logger.error(f"Failed to process JD PDF {jd_file.filename}: {e}")
            raise HTTPException(status_code=400, detail=f"Could not process JD PDF file: {e}")
        finally:
            if 'tmp_path' in locals() and os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    if not effective_jd.strip() or not skill_list:
        raise HTTPException(status_code=422, detail="Job Description (from text or PDF) and Required Skills are required.")
    
    file_data_list = [(file.filename, await file.read()) for file in files]
    if not file_data_list:
        raise HTTPException(status_code=400, detail="No valid resume files could be read.")
    
    initial_state = AgentState(
        job_title=job_title,
        job_description=effective_jd,
        required_skills=skill_list,
        file_data_list=file_data_list,
        top_n=top_n,
        groq_api_key=groq_api_key,
        section_documents=[],
        processed_resumes={},
        semantic_candidates=[],
        ranked_candidates=[]
    )
    
    try:
        final_state = await app_graph.ainvoke(initial_state)

        top_candidates = []
        for i, c in enumerate(final_state["ranked_candidates"][:top_n]):
            top_candidates.append(RankedResume(
                rank=i + 1,
                filename=c['filename'],
                semantic_score=c['score'],
                skill_match_count=c['skill_match_count'],
                total_skills_required=len(skill_list),
                matched_skills=c['matched_skills'],
                missing_skills=c['missing_skills'],
                content_preview=final_state["processed_resumes"][c['filename']]['full_text'][:400] + "...",
                key_matching_chunk=c.get('key_matching_chunk', "N/A"),
                ai_verdict=c.get('ai_verdict', "N/A"),
                ai_feedback=c.get('ai_feedback', "N/A")
            ))

        processing_time = (datetime.now() - start_time).total_seconds()

        return RankingResult(
            job_title=job_title,
            total_resumes_processed=len(final_state["processed_resumes"]),
            top_candidates=top_candidates,
            processing_time=processing_time,
        )
    except Exception as e:
        logger.error(f"An unexpected error occurred during workflow execution: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An internal error occurred: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

